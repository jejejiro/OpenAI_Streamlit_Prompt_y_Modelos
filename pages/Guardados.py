import streamlit as st
from docx import Document
from io import BytesIO



# Crea el documento word
def crearWord(chats):
    doc = Document() # crea el documento
    doc.add_heading('Chats Descargados', 0) # le agrega un titulo
    doc.add_paragraph(chats) # Agrega los datos
    buffer = BytesIO()  # salva el documento en el buffer y lo retorna
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():

    st.write("Chats Guardados")
    for dato in st.session_state.historial:
        st.write(dato)

    # Hace la descarga del documento word
    st.download_button(
            'Descargar Word',
            data=crearWord(dato),
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

if __name__ == "__main__":

    st.set_page_config(page_title="Prueba tus Modelos y Prompt", layout="wide", page_icon="lion-head-logo.png")
    main()